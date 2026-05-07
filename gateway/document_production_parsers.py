"""Gateway Document Production Parsers - optional PDF/Office parser set.

Purpose: register real parser implementations for PDF, DOCX, XLSX, and PPTX
inputs when their runtime libraries are installed.
Governance scope: deterministic extraction, dependency availability checks,
structured table output, parser health reports, and fail-closed parse errors.
Dependencies: mcoi artifact parser contracts, optional pypdf/python-docx/
openpyxl libraries, and stdlib Open XML parsing for PPTX.
Invariants:
  - Parsers are registered only when their backing library can be imported.
  - Each parser produces normalized output through the shared parser contract.
  - Parser output never performs external send, sign, or submit effects.
  - Missing optional dependencies are explicit and never silently simulated.
"""

from __future__ import annotations

from abc import abstractmethod
from datetime import datetime, timezone
from hashlib import sha256
import importlib.util
import io
from pathlib import Path
from typing import Any, Mapping
from xml.etree import ElementTree
import zipfile

from mcoi_runtime.contracts.artifact_parser import (
    ArtifactParserDescriptor,
    NormalizedParseOutput,
    ParseCapability,
    ParseOutputKind,
    ParserCapabilityLevel,
    ParserCapabilityManifest,
    ParserFamily,
    ParserHealthReport,
    ParserStatus,
)
from mcoi_runtime.core.artifact_parsers import ArtifactParser, ArtifactParserRegistry
from mcoi_runtime.core.invariants import stable_identifier


def register_optional_production_parsers(registry: ArtifactParserRegistry) -> int:
    """Register available production parsers and return the registered count."""
    registered_count = 0
    for parser in (
        ProductionPDFParser(),
        ProductionDOCXParser(),
        ProductionXLSXParser(),
        ProductionPPTXParser(),
    ):
        if parser.available:
            registry.register(parser)
            registered_count += 1
    return registered_count


class ProductionParser(ArtifactParser):
    """Shared production-parser contract implementation."""

    parser_name: str
    parser_family: ParserFamily
    parser_extensions: tuple[str, ...]
    parser_mime_types: tuple[str, ...]
    parser_output_kind: ParseOutputKind = ParseOutputKind.TEXT
    parser_libraries: tuple[str, ...]
    max_size_bytes: int = 104_857_600

    def __init__(self, parser_id: str) -> None:
        self._parser_id = parser_id
        self._created_at = _now_iso()
        self._parsed_count = 0
        self._failed_count = 0

    @property
    def available(self) -> bool:
        return all(importlib.util.find_spec(library) is not None for library in self.parser_libraries)

    def parser_id(self) -> str:
        return self._parser_id

    def family(self) -> ParserFamily:
        return self.parser_family

    def descriptor(self) -> ArtifactParserDescriptor:
        return ArtifactParserDescriptor(
            parser_id=self._parser_id,
            name=self.parser_name,
            family=self.parser_family,
            status=ParserStatus.AVAILABLE if self.available else ParserStatus.UNAVAILABLE,
            version="1.0.0",
            manifest_id=_manifest_id(self._parser_id),
            tags=("production",),
            created_at=self._created_at,
            metadata={"libraries": self.parser_libraries},
        )

    def manifest(self) -> ParserCapabilityManifest:
        return ParserCapabilityManifest(
            manifest_id=_manifest_id(self._parser_id),
            parser_id=self._parser_id,
            family=self.parser_family,
            capabilities=(
                ParseCapability(
                    format_name=self.parser_family.value,
                    extensions=self.parser_extensions,
                    mime_types=self.parser_mime_types,
                    capability_level=ParserCapabilityLevel.FULL_CONTENT,
                    max_size_bytes=self.max_size_bytes,
                    output_kinds=(self.parser_output_kind.value,),
                ),
            ),
            reliability_score=0.9,
            created_at=self._created_at,
            metadata={"dependency_checked": True},
        )

    def can_parse(self, filename: str, mime_type: str, size_bytes: int) -> bool:
        if not self.available or size_bytes > self.max_size_bytes:
            return False
        suffix = Path(filename).suffix.lower()
        return suffix in self.parser_extensions or mime_type in self.parser_mime_types

    def parse(
        self,
        artifact_id: str,
        filename: str,
        content: bytes,
    ) -> NormalizedParseOutput:
        if not self.available:
            self._failed_count += 1
            raise RuntimeError(f"parser dependency unavailable: {self._parser_id}")
        try:
            output = self._parse_available(artifact_id=artifact_id, filename=filename, content=content)
        except Exception:
            self._failed_count += 1
            raise
        self._parsed_count += 1
        return output

    def health_check(self) -> ParserHealthReport:
        return ParserHealthReport(
            report_id=stable_identifier("parser-health", {"parser_id": self._parser_id, "ts": _now_iso()}),
            parser_id=self._parser_id,
            status=ParserStatus.AVAILABLE if self.available else ParserStatus.UNAVAILABLE,
            reliability_score=0.9 if self.available else 0.0,
            artifacts_parsed=self._parsed_count,
            artifacts_failed=self._failed_count,
            avg_parse_ms=0.0,
            active_failure_modes=() if self.available else ("dependency_unavailable",),
            reported_at=_now_iso(),
        )

    @abstractmethod
    def _parse_available(
        self,
        *,
        artifact_id: str,
        filename: str,
        content: bytes,
    ) -> NormalizedParseOutput:
        """Parse content after dependency availability has been proven."""


class ProductionPDFParser(ProductionParser):
    parser_name = "Production PDF Parser"
    parser_family = ParserFamily.DOCUMENT
    parser_extensions = (".pdf",)
    parser_mime_types = ("application/pdf",)
    parser_libraries = ("pypdf",)

    def __init__(self) -> None:
        super().__init__("production-pdf")

    def _parse_available(
        self,
        *,
        artifact_id: str,
        filename: str,
        content: bytes,
    ) -> NormalizedParseOutput:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = tuple(page.extract_text() or "" for page in reader.pages)
        metadata = {str(key): str(value) for key, value in (reader.metadata or {}).items()}
        return _normalized_output(
            parser_id=self._parser_id,
            artifact_id=artifact_id,
            family=self.parser_family,
            output_kind=ParseOutputKind.TEXT,
            content=content,
            text="\n".join(pages),
            page_count=len(pages),
            extracted_metadata={"filename": filename, **metadata},
        )


class ProductionDOCXParser(ProductionParser):
    parser_name = "Production DOCX Parser"
    parser_family = ParserFamily.DOCUMENT
    parser_extensions = (".docx",)
    parser_mime_types = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)
    parser_libraries = ("docx",)

    def __init__(self) -> None:
        super().__init__("production-docx")

    def _parse_available(
        self,
        *,
        artifact_id: str,
        filename: str,
        content: bytes,
    ) -> NormalizedParseOutput:
        from docx import Document

        document = Document(io.BytesIO(content))
        paragraphs = tuple(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        tables = tuple(_docx_table(table) for table in document.tables)
        return _normalized_output(
            parser_id=self._parser_id,
            artifact_id=artifact_id,
            family=self.parser_family,
            output_kind=ParseOutputKind.TEXT,
            content=content,
            text="\n".join(paragraphs),
            tables=tables,
            page_count=1,
            extracted_metadata={"filename": filename},
        )


class ProductionXLSXParser(ProductionParser):
    parser_name = "Production XLSX Parser"
    parser_family = ParserFamily.SPREADSHEET
    parser_extensions = (".xlsx", ".xlsm")
    parser_mime_types = ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",)
    parser_output_kind = ParseOutputKind.TABLE
    parser_libraries = ("openpyxl",)

    def __init__(self) -> None:
        super().__init__("production-xlsx")

    def _parse_available(
        self,
        *,
        artifact_id: str,
        filename: str,
        content: bytes,
    ) -> NormalizedParseOutput:
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        tables = tuple(_worksheet_table(sheet) for sheet in workbook.worksheets)
        text = "\n".join(_table_to_text(table) for table in tables)
        return _normalized_output(
            parser_id=self._parser_id,
            artifact_id=artifact_id,
            family=self.parser_family,
            output_kind=ParseOutputKind.TABLE,
            content=content,
            text=text,
            tables=tables,
            page_count=len(tables),
            extracted_metadata={"filename": filename, "sheet_count": str(len(tables))},
        )


class ProductionPPTXParser(ProductionParser):
    parser_name = "Production PPTX Parser"
    parser_family = ParserFamily.PRESENTATION
    parser_extensions = (".pptx",)
    parser_mime_types = ("application/vnd.openxmlformats-officedocument.presentationml.presentation",)
    parser_libraries = ("zipfile", "xml.etree.ElementTree")

    def __init__(self) -> None:
        super().__init__("production-pptx")

    def _parse_available(
        self,
        *,
        artifact_id: str,
        filename: str,
        content: bytes,
    ) -> NormalizedParseOutput:
        slide_text = _pptx_slide_texts(content)
        return _normalized_output(
            parser_id=self._parser_id,
            artifact_id=artifact_id,
            family=self.parser_family,
            output_kind=ParseOutputKind.TEXT,
            content=content,
            text="\n\n".join(slide_text),
            page_count=len(slide_text),
            extracted_metadata={"filename": filename, "slide_count": str(len(slide_text))},
        )


def _normalized_output(
    *,
    parser_id: str,
    artifact_id: str,
    family: ParserFamily,
    output_kind: ParseOutputKind,
    content: bytes,
    text: str,
    tables: tuple[Mapping[str, Any], ...] = (),
    page_count: int = 0,
    extracted_metadata: Mapping[str, Any] | None = None,
) -> NormalizedParseOutput:
    return NormalizedParseOutput(
        output_id=stable_identifier(
            "parse-output",
            {
                "parser_id": parser_id,
                "artifact_id": artifact_id,
                "content_hash": sha256(content).hexdigest(),
            },
        ),
        parser_id=parser_id,
        artifact_id=artifact_id,
        family=family,
        output_kind=output_kind,
        text_content=text,
        structured_data={"content_hash": sha256(content).hexdigest()},
        tables=tables,
        page_count=page_count,
        word_count=len(text.split()),
        has_images=False,
        has_tables=bool(tables),
        language_hint="",
        extracted_metadata=dict(extracted_metadata or {}),
        parsed_at=_now_iso(),
    )


def _docx_table(table: Any) -> Mapping[str, Any]:
    rows = tuple(tuple(cell.text for cell in row.cells) for row in table.rows)
    headers = rows[0] if rows else ()
    return {
        "headers": headers,
        "rows": rows[1:] if rows else (),
        "row_count": max(len(rows) - 1, 0),
    }


def _worksheet_table(sheet: Any) -> Mapping[str, Any]:
    rows = tuple(
        tuple("" if value is None else str(value) for value in row)
        for row in sheet.iter_rows(values_only=True)
    )
    headers = rows[0] if rows else ()
    return {
        "sheet_name": sheet.title,
        "headers": headers,
        "rows": rows[1:] if rows else (),
        "row_count": max(len(rows) - 1, 0),
    }


def _table_to_text(table: Mapping[str, Any]) -> str:
    lines = ["\t".join(str(value) for value in table.get("headers", ()))]
    lines.extend("\t".join(str(value) for value in row) for row in table.get("rows", ()))
    return "\n".join(line for line in lines if line)


def _pptx_slide_texts(content: bytes) -> tuple[str, ...]:
    slide_texts: list[str] = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        slide_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for slide_name in slide_names:
            root = ElementTree.fromstring(archive.read(slide_name))
            texts = tuple(
                element.text or ""
                for element in root.iter()
                if _xml_local_name(element.tag) == "t" and (element.text or "").strip()
            )
            slide_texts.append("\n".join(texts))
    return tuple(slide_texts)


def _xml_local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _manifest_id(parser_id: str) -> str:
    return stable_identifier("parser-manifest", {"parser_id": parser_id})


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()
